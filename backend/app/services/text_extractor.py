"""Extract text from documents."""
import os
import re
from pathlib import Path

# Repeated administrative/legal boilerplate found in this corpus's official PDFs — the
# BSrE/BSSN e-signature notice appears near-verbatim in 96.7% of approved chunks (found via
# direct DB audit during evaluation debugging). It carries zero topical content, but because
# it's identical across nearly every chunk it dilutes both dense embeddings (every chunk's
# pooled vector spends part of its "budget" representing this constant sentence instead of
# the chunk's actual content) and, to a lesser extent, BM25 term weighting. Stripped here —
# once, at extraction time, before chunking — so it never enters chunk boundaries, summaries,
# embeddings, or the BM25 index at all. `\s+` between words tolerates the line-break/whitespace
# variants seen in different source PDFs; case-insensitive since capitalization varies.
_BOILERPLATE_PATTERNS = [
    re.compile(
        r"Dokumen\s+ini\s+telah\s+ditandatangani\s+secara\s+elektronik\s+menggunakan\s+"
        r"sertifikat\s+elektronik\s+yang\s+diterbitkan\s+oleh\s+Balai\s+Besar\s+Sertifikasi\s+"
        r"Elektronik\s*\(BSrE\)\s*,?\s*Badan\s+Siber\s+dan\s+Sandi\s+Negara\s*\(BSSN\)\s*\.?",
        re.IGNORECASE,
    ),
    # `${nomor_naskah}` / `${nama_pengirim}` / `${nip_pengirim}` / `${ttd_pengirim}` — unfilled
    # e-signature template placeholder tokens found live on the SK-preamble pages of the
    # ground-truth Pedoman PDF (structure-aware document extraction plan, Phase 0 audit).
    # Unlike the BSrE/BSSN sentence above these were never stripped before, so they leaked
    # into chunk text/embeddings/summaries verbatim as `${...}`.
    re.compile(r"\$\{[a-z_]+\}"),
]


class TextExtractor:
    """Extract text from various document formats."""

    @staticmethod
    def _clean_boilerplate(text: str) -> str:
        """Strip known repeated non-content administrative text, then collapse the
        whitespace gaps left behind so chunking doesn't split on an empty run."""
        for pattern in _BOILERPLATE_PATTERNS:
            text = pattern.sub(" ", text)
        return re.sub(r"[ \t]{2,}", " ", text).strip()

    @staticmethod
    async def extract_from_file(filepath: str) -> str:
        """Extract text from file based on extension."""
        if not os.path.exists(filepath):
            return ""

        ext = Path(filepath).suffix.lower()

        if ext == ".pdf":
            text = await TextExtractor._extract_pdf(filepath)
        elif ext == ".docx":
            text = await TextExtractor._extract_docx(filepath)
        elif ext in [".html", ".htm", ".txt"]:
            text = await TextExtractor._extract_text(filepath)
        else:
            return ""

        return TextExtractor._clean_boilerplate(text) if text else text

    @staticmethod
    async def _extract_text(filepath: str) -> str:
        """Extract from plain text or HTML files."""
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            
            # Simple HTML tag stripping for HTML files
            import re
            content = re.sub(r"<[^>]+>", "", content)
            
            # Clean up whitespace
            content = re.sub(r"\s+", " ", content).strip()
            
            return content
        except Exception:
            return ""

    @staticmethod
    def _render_table_markdown(rows: list) -> str:
        """Render PyMuPDF's `Table.extract()` output (list of rows, each a list of cell
        strings/None) as a pipe-delimited markdown table, so every cell's data survives
        chunking/summarization as clearly structured text instead of being lost or scrambled."""
        lines = []
        for i, row in enumerate(rows):
            cells = [str(c).strip() if c is not None else "" for c in row]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join("---" for _ in cells) + " |")
        return "\n".join(lines)

    @staticmethod
    def _extract_page_text(page, page_num: int) -> tuple[str, int]:
        """Shared per-page extraction: plain text + every detected table rendered as a
        markdown block, appended. Returns (page_text, table_count). Factored out of
        `_extract_pdf` so `extract_pages_from_pdf` below can reuse the exact same per-page
        logic instead of re-deriving it — the structure-aware document extraction plan needs
        a page-indexed view (page 20 is a divider, page 26 is image-only) that is impossible
        once every page's text is concatenated into one flat string, which is what
        `_extract_pdf`'s old inline loop did.

        Table detection is via `page.find_tables()` (table-structure detection, not OCR — the
        same API `image_extractor.py::extract_table_screenshots_from_pdf` already uses to
        screenshot tables for the vision-review pipeline). `page.get_text()` alone extracts a
        table's cells as whatever run of words its default reading order produces, which can
        merge/scramble columns — the markdown block guarantees a structured, complete copy of
        every row reaches `original_text` regardless. Deliberately additive (the table's text
        still also appears inline via get_text()) rather than trying to excise the table's
        bbox from the flowing text, which risks dropping adjacent non-table content on an
        imprecise cut.
        """
        text = page.get_text()
        table_count = 0

        try:
            found = page.find_tables()
        except Exception:
            # Table detection failing must not take down text extraction for this page —
            # the plain text above is still kept.
            return text, table_count

        for table in found.tables:
            try:
                rows = table.extract()
            except Exception:
                continue
            if not rows:
                continue
            table_count += 1
            markdown = TextExtractor._render_table_markdown(rows)
            text += f"\n\n[Tabel halaman {page_num}]\n{markdown}\n"

        return text, table_count

    @staticmethod
    async def _extract_pdf(filepath: str) -> str:
        """Extract text from PDF using PyMuPDF (flat string, existing behavior — every
        non-zoned ingestion path keeps using this unchanged)."""
        try:
            import fitz
            text = ""
            with fitz.open(filepath) as doc:
                for page_num, page in enumerate(doc, 1):
                    page_text, _ = TextExtractor._extract_page_text(page, page_num)
                    text += page_text
            return text.strip()
        except Exception:
            return ""

    @staticmethod
    async def extract_pages_from_pdf(filepath: str) -> list[dict]:
        """Page-indexed PDF extraction: `[{page, text, image_count, table_count}, ...]`.

        Prerequisite for DocumentStructureAgent (structure-aware document extraction plan) —
        it needs to tell "page 20 is a bare divider" from "page 21 is a form" and "page 26 is
        image-only", none of which survives `_extract_pdf`'s flat-string concatenation. Each
        page's text is boilerplate-cleaned individually (same `_clean_boilerplate` the flat
        path uses) so downstream zone detection never sees the BSrE/BSSN footer or
        `${...}` placeholder tokens either.
        """
        try:
            import fitz
        except ImportError:
            return []

        pages: list[dict] = []
        try:
            with fitz.open(filepath) as doc:
                for page_num, page in enumerate(doc, 1):
                    page_text, table_count = TextExtractor._extract_page_text(page, page_num)
                    try:
                        image_count = len(page.get_images(full=True))
                    except Exception:
                        image_count = 0
                    pages.append({
                        "page": page_num,
                        "text": TextExtractor._clean_boilerplate(page_text),
                        "image_count": image_count,
                        "table_count": table_count,
                    })
            return pages
        except Exception:
            return []

    @staticmethod
    def _iter_block_items(doc):
        """Yield each paragraph and table in `doc.element.body` in true document order.
        `doc.paragraphs`/`doc.tables` on their own are each flattened, order-independent
        collections (by design, per python-docx) — walking the raw XML body children is the
        standard recipe for reconstructing the original reading order of mixed content."""
        from docx.oxml.ns import qn
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        for child in doc.element.body.iterchildren():
            if child.tag == qn("w:p"):
                yield Paragraph(child, doc)
            elif child.tag == qn("w:tbl"):
                yield Table(child, doc)

    @staticmethod
    async def _extract_docx(filepath: str) -> str:
        """Extract text from DOCX, including table content.

        `doc.paragraphs` (the previous implementation) excludes anything inside `doc.tables`
        by design — DOCX tables were previously 100% unextracted, not merely lossy. Walks the
        document body in order so table content lands at its correct position relative to the
        surrounding narrative, rendering each table as the same markdown block `_extract_pdf`
        uses so every cell survives chunking/summarization as clear, structured text.
        """
        try:
            from docx import Document
            from docx.table import Table as DocxTable

            doc = Document(filepath)
            parts = []
            for block in TextExtractor._iter_block_items(doc):
                if isinstance(block, DocxTable):
                    rows = [[cell.text for cell in row.cells] for row in block.rows]
                    if rows:
                        parts.append(f"[Tabel]\n{TextExtractor._render_table_markdown(rows)}")
                elif block.text.strip():
                    parts.append(block.text)
            return "\n".join(parts).strip()
        except ImportError:
            # Fallback if python-docx not installed
            return ""
        except Exception:
            return ""

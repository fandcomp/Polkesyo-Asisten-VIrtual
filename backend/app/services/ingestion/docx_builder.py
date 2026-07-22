"""Build standalone .docx files for form/attachment zone pages (structure-aware document
extraction plan) — the shared renderer `FormExtractionAgent` and `VisionFormConversionAgent`
both call.

First real WRITE usage of python-docx in this codebase (already a dependency, per
pyproject.toml) — `text_extractor.py::_extract_docx` only ever reads .docx files.

Verbatim by design: CLAUDE.md §21's admin-approval workflow and the user's explicit
requirement ("isinya table persis sesuai di file dokumen pedoman itu") mean this never
paraphrases or summarizes anything — it copies field labels/values and table cells exactly
as extracted, character-for-character.
"""
from pathlib import Path
from typing import Any

from docx import Document


def _add_field(doc: Document, label: str, value: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(label)
    run.bold = True
    if value:
        paragraph.add_run(f": {value}")


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    n_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Table Grid"
    for row in rows:
        cells = table.add_row().cells
        for idx in range(n_cols):
            cells[idx].text = row[idx] if idx < len(row) else ""


def build_form_docx(
    title: str,
    fields: list[tuple[str, str]],
    tables: list[list[list[str]]],
    output_path: str,
    items: list[tuple[str, Any]] | None = None,
) -> str:
    """Render a form verbatim into a new .docx at `output_path`. Returns `output_path`.

    `items`, when given, is the ordered sequence of `("field", (label, value))` /
    `("table", rows)` entries exactly as they appeared in the source document — this is the
    faithful, preferred path (2026-07-22: the source document rarely lists every field before
    every table; splitting them into two separately-rendered groups silently reordered content
    relative to the original, which is the opposite of the verbatim requirement below).

    `fields`/`tables` (rendered as two separate blocks, fields first) remain supported for
    `VisionFormConversionAgent`, whose vision-model extraction returns them as two flat
    buckets with no original-order signal between them to preserve in the first place — pass
    `items=None` (the default) for that caller."""
    doc = Document()
    doc.add_heading(title or "Formulir", level=1)

    if items is not None:
        for kind, payload in items:
            if kind == "field":
                label, value = payload
                _add_field(doc, label, value)
            elif kind == "table":
                _add_table(doc, payload)
    else:
        for label, value in fields:
            _add_field(doc, label, value)
        for rows in tables:
            _add_table(doc, rows)

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    return output_path

"""Unit tests for docx_builder.build_form_docx (structure-aware document extraction plan).

Verifies the verbatim requirement: whatever fields/table rows go in must come back out of the
generated .docx exactly, with no paraphrasing or reformatting.
"""
import tempfile
from pathlib import Path

from docx import Document

from app.services.ingestion.docx_builder import build_form_docx


class TestBuildFormDocx:
    def test_renders_title_as_heading(self):
        with tempfile.TemporaryDirectory() as tmp:
            output_path = str(Path(tmp) / "form.docx")
            build_form_docx("SURAT PERNYATAAN", [], [], output_path)

            doc = Document(output_path)
            assert doc.paragraphs[0].text == "SURAT PERNYATAAN"

    def test_renders_fields_verbatim(self):
        fields = [("Nama", "Budi Santoso"), ("NIK", "3471xxxxxxxxxxxx"), ("Alamat", "")]
        with tempfile.TemporaryDirectory() as tmp:
            output_path = str(Path(tmp) / "form.docx")
            build_form_docx("Formulir", fields, [], output_path)

            doc = Document(output_path)
            body_text = [p.text for p in doc.paragraphs]
            assert "Nama: Budi Santoso" in body_text
            assert "NIK: 3471xxxxxxxxxxxx" in body_text
            # Empty value fields render the label alone (no trailing ": ").
            assert "Alamat" in body_text
            assert "Alamat: " not in body_text

    def test_renders_table_rows_verbatim(self):
        rows = [
            ["No", "Nama", "Status"],
            ["1", "2410001", "Lulus"],
            ["2", "2410002", "Tidak Lulus"],
        ]
        with tempfile.TemporaryDirectory() as tmp:
            output_path = str(Path(tmp) / "form.docx")
            build_form_docx("Hasil Seleksi", [], [rows], output_path)

            doc = Document(output_path)
            assert len(doc.tables) == 1
            table = doc.tables[0]
            assert [c.text for c in table.rows[0].cells] == ["No", "Nama", "Status"]
            assert [c.text for c in table.rows[1].cells] == ["1", "2410001", "Lulus"]
            assert [c.text for c in table.rows[2].cells] == ["2", "2410002", "Tidak Lulus"]

    def test_handles_ragged_table_rows(self):
        """A row shorter than the widest row must not crash — missing cells render empty."""
        rows = [["A", "B", "C"], ["1", "2"]]
        with tempfile.TemporaryDirectory() as tmp:
            output_path = str(Path(tmp) / "form.docx")
            build_form_docx("Tabel", [], [rows], output_path)

            doc = Document(output_path)
            table = doc.tables[0]
            assert [c.text for c in table.rows[1].cells] == ["1", "2", ""]

    def test_creates_parent_directories(self):
        with tempfile.TemporaryDirectory() as tmp:
            output_path = str(Path(tmp) / "nested" / "dir" / "form.docx")
            result = build_form_docx("Formulir", [("A", "B")], [], output_path)

            assert result == output_path
            assert Path(output_path).is_file()
